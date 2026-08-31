"""In-memory .docx builders -- separate from services/artifact_service.py's
SubmissionArtifactService, which writes to a local file path; these return
raw bytes for callers (like LetterService.upload_redacted_to_sponsor_onedrive)
that hand the result straight to a StorageService instead of the local disk."""

from io import BytesIO

from docx import Document


def build_blank_reply_docx() -> bytes:
    """The blank reply doc dropped into each exchangeX upload -- a single
    placeholder line, not empty and not a longer template (Rey's call,
    22Aug2026: "I usually type 'Respond here'")."""
    doc = Document()
    doc.add_paragraph("Respond here")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_translation_review_docx(pages: list, personal_use: bool = False) -> bytes:
    """
    Draft translation doc for Letter Mgt's Spanish-language workflow (added
    31Aug2026) -- meant to be downloaded and sent to a bilingual reviewer
    OUTSIDE the app (email/text, however Rey already communicates with
    them; no in-app reviewer role exists), corrected, and the result
    re-uploaded. Explicitly labeled as a draft needing review, not a
    finished document, so it's never mistaken for the real thing if it
    ends up somewhere it shouldn't.

    `pages` is a list of dicts, one per letter page, each with
    `original_text`, `translation`, and optionally `detected_language`.

    `personal_use` (added 31Aug2026, standalone Translate tool): set when
    this is Rey reading a letter from his own sponsee rather than a
    sponsor-facing exchange -- there's no reviewer and nothing is ever
    uploaded, so the "needs review before sending to the sponsor" framing
    is wrong and was confusing him. Same doc structure, different heading.
    """
    doc = Document()
    if personal_use:
        doc.add_heading("LETTER TRANSLATION (for your own reading)", level=1)
        doc.add_paragraph(
            "Machine-translated (local, offline model) from the original letter. "
            "Not reviewed -- may contain errors."
        )
    else:
        doc.add_heading("DRAFT TRANSLATION -- NEEDS BILINGUAL REVIEW BEFORE USE", level=1)
        doc.add_paragraph(
            "Machine-translated (local, offline model) from the original letter. "
            "Review and correct before sending to the sponsor."
        )
    for i, page in enumerate(pages, start=1):
        doc.add_heading(f"Page {i}", level=2)
        if page.get("detected_language"):
            doc.add_paragraph(f"Detected language: {page['detected_language']}")
        doc.add_heading("Original", level=3)
        doc.add_paragraph(page.get("original_text") or "(none)")
        doc.add_heading("Draft English Translation", level=3)
        doc.add_paragraph(page.get("translation") or "(none)")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
